# cd desktop
# uvicorn Resume_parser:app --reload

import re
from collections import defaultdict
from docx import Document
import os
from PyPDF2 import PdfReader
from fastapi import FastAPI, UploadFile
import spacy
from spacy.matcher import Matcher

# spaCy model
nlp = spacy.load('en_core_web_sm')
matcher = Matcher(nlp.vocab)

# FastAPI instance
app = FastAPI()

# Root endpoint for testing connectivity
@app.get("/")
def read_root():
    return {"message": "Welcome to the Resume Parsing API!"}

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    text = ""
    try:
        pdf_reader = PdfReader(file_path)
        for page in pdf_reader.pages:
            text += page.extract_text() or "" 
        if not text:
            print("Warning: No text extracted from PDF.")
    except Exception as e:
        print(f"Error reading PDF: {e}")
    
    return text

# Text cleaning function
def clean_text(text):
    # Replace newlines with spaces
    text = text.replace("\n", " ")
    
    # Remove non-alphanumeric characters (except spaces)
    text = re.sub(r"[^a-zA-Z0-9\s]", " ", text)  # Replace everything except letters, numbers, and spaces with a space

    # Convert text to lowercase
    text = text.lower()

    # Remove extra spaces
    text = re.sub(r'\s+', ' ', text).strip()

    return text

def extract_name(text):
    nlp_text = nlp(text)
  
    # First and Last name are always Proper Nouns
    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    matcher.add('NAME', [pattern], on_match=None)
    matches = matcher(nlp_text)
  
    for match_id, start, end in matches:  # Extract match_id, start, and end
        span = nlp_text[start:end]
        return span.text

# Function to parse resume text
def parse_resume(text):
    sections = defaultdict(list)

    # regular expressions for extracting email, and phone
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b')  #  email regex
    phone_pattern = re.compile(r'(\+?\d{1,2}\s?)?(\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4})')  # phone pattern (11 digits limit)

    name = False
    email = False
    phone = False

    # Extract name using spaCy
    name = extract_name(text)

    # Detect email and phone by scanning lines
    lines = text.splitlines()

    # Check for email and phone numbers
    for line in lines:
        line = line.strip()

        # Check for email
        if not email and re.search(email_pattern, line):  # Check if email is present
            email = re.search(email_pattern, line).group()
            continue  # Skip further checks if email is found

        # Check for phone number (limit to 11 digits)
        if not phone and re.search(phone_pattern, line):  # Check if phone is present
            phone = re.search(phone_pattern, line).group()
            continue  # Skip further checks if phone is found

    # Add name, email, and phone to the sections
    sections["Name"].append(name if name else "Not Available")
    sections["Email"].append(email if email else "Not Available")
    sections["Phone"].append(phone if phone else "Not Available")

    # Clean the text before processing
    text = clean_text(text)

    # Section keywords
    section_keywords = {
        "Education": ["education", "academic", "degree", "graduation", "qualification", "university", "college"],
        "Work Experience": ["experience", "work", "employment", "job", "internship"],
        "Skills": ["skills", "technical", "tools", "expertise", "proficiencies"],
        "Certifications": ["certifications", "courses", "training", "workshops"],
        "Achievements": ["accomplishments", "awards", "honors", "achievements"],
        "Projects": ["projects", "applications", "development", "personal projects"],
        "Languages": ["languages", "language proficiency"],
        "Interests": ["interests", "hobbies", "passions", "leisure"],
        "Summary": ["summary", "executive profile", "professional profile", "personal profile", "work background"],
        "Publications": ["publications", "publication"],
        "Other Activities": ["other activities", "volunteer", "community"],
        "Objective": ["objective", "career goal", "career objective"]
    }

    current_section = None
    for line in lines:
        line = line.strip()
        if not line:
            continue  # Skip empty lines

        section_matched = False
        for section, keywords in section_keywords.items():
            if any(re.match(f"^{keyword}", line, re.IGNORECASE) for keyword in keywords):
                current_section = section
                sections[current_section].append(line)
                section_matched = True
                break

        if not section_matched and current_section:
            sections[current_section].append(line)


    # Ensure every section has at least "Not Available" if empty
    for section in sections:
        if not sections[section]:
            sections[section].append("Not Available")

    return sections

# Function to save parsed data to DOCX
def save_to_docx(parsed_sections, output_path):
    print(f"Saving parsed sections to DOCX: {output_path}")
    document = Document()
    document.add_heading("Parsed Resume", level=1)

    # Create a table 
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Section"
    header_cells[1].text = "Details"

    # Add Name, Email, and Phone first
    for key in ["Name", "Email", "Phone"]:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = "\n".join(parsed_sections[key]) if parsed_sections[key] else "Not Available"

    # Populate the table with other sections
    for section, content in parsed_sections.items():
        if section not in ["Name", "Email", "Phone"]:  # Skip name, email, phone as they were added earlier
            row_cells = table.add_row().cells
            row_cells[0].text = section
            row_cells[1].text = "\n".join(content) if content else "Not Available"

    document.save(output_path)

# API endpoint to upload and parse resume
@app.post("/parse-resume/")
async def parse_resume_endpoint(file: UploadFile):
    # Save uploaded file to a temporary path
    temp_path = f"./{file.filename}"
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    # Extract text from the PDF
    text = extract_text_from_pdf(temp_path)

    # Parse the resume text
    parsed_data = parse_resume(text)

    # Save parsed data to DOCX
    output_path = f"./{file.filename.split('.')[0]}_parsed.docx"
    save_to_docx(parsed_data, output_path)

    # Clean up temporary file (optional)
    os.remove(temp_path)

    # Return a success message
    return {"message": "Resume parsed successfully!", "docx_file": output_path}
